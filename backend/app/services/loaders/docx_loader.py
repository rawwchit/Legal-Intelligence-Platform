from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document

from app.services.loaders.base import BaseLoader


class DOCXLoader(BaseLoader):
    """Load Microsoft Word (.docx) documents."""

    def load(self, file_path: Path) -> list[Document]:
        try:
            document = DocxDocument(file_path)

            paragraphs = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]

            return [
                Document(
                    page_content="\n".join(paragraphs),
                    metadata={
                        "source": str(file_path),
                        "file_type": "docx",
                    },
                )
            ]

        except Exception as e:
            raise ValueError(
                f"Failed to load DOCX document: {file_path}"
            ) from e